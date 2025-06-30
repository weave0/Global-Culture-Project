def read_batch_text_file(path):
    """Read and print the content of a batch text file for verification."""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            batch_content = f.read()
            print("File read successfully.")
            print(batch_content[:500])  # Print the first 500 characters for verification
    except FileNotFoundError:
        print(f"Error: {path} not found.")
        return


def extract_docx_to_excel(docx_path, output_path):
    """Extract paragraphs and tables from a DOCX and save to Excel."""
    import pandas as pd
    import docx
    from datetime import datetime

    doc = docx.Document(docx_path)
    content_blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content_blocks.append({"Type": "Paragraph", "Content": text})

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                content_blocks.append({"Type": "TableRow", "Content": row_text})

    df = pd.DataFrame(content_blocks)
    df["Source File"] = docx_path
    df["Extracted At"] = datetime.utcnow().isoformat()

    df.to_excel(output_path, index=False)
    print(f"Extracted {len(df)} content blocks from DOCX and saved to Excel.")


if __name__ == "__main__":
    # Example usage for batch text file
    read_batch_text_file("d:\\Global Culture Project\\Global Culture Alignment\\Global_Culture_Profiles_Batch_1.txt")

    # Example usage for DOCX extraction
    # extract_docx_to_excel(
    #     r"D:\Global Culture Project\PGLS Files\Global Culture Guide\Initial Outputs\PGLS Global Culture Guide v1bw62726.docx",
    #     r"D:\Global Culture Project\Global Culture Project CORE.xlsx"
    # )
